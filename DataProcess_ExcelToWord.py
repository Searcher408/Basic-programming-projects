# coding:utf-8
# 读excel文件，写word文件
import sys
import importlib
import xlrd
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def readExcel(all_content):
    worksheet = xlrd.open_workbook(u'E:\\Yoc\\Course\\2021.xlsx')
    sheet_names= worksheet.sheet_names()
    # print(sheet_names) # ['Sheet1', 'Sheet2', 'Sheet3']

    sheet = worksheet.sheet_by_name('Sheet1')
    rows = sheet.nrows # 获取行数
    cols = sheet.ncols # 获取列数

    #print(rows) # 198
    #print(cols) # 10

    # 输出每一行数据
    # all_content = []
    row_content = []
    for i in range(2, rows) :
        row_content = []
        for j in range(0, 6) :
            cell = sheet.cell_value(i, j) # 取第i行第j列数据 
            cell = str(cell) 
            if j == 5:     
                if cell and cell[-2] == '.':
                    cell = cell[:-2]
                #print(cell)
            row_content.append(cell)
        row_content.append(sheet.cell_value(i, 7))
        all_content.append(row_content)
    print(all_content)

def writeTxt(all_content):
    with open('E:\\Yoc\\Course\\2021.txt', 'w', encoding='utf-8') as f:
        for lines in all_content:
            if lines[0]:
                f.writelines('\n\n'+lines[0]+'校友会\n\n')
            f.writelines('职务：'+lines[1]+'\n')
            f.writelines('姓名：'+lines[2]+'\n')
            f.writelines('毕业学院：'+lines[3]+'\n')
            f.writelines('毕业时间：'+lines[4]+'\n')
            f.writelines('手机：'+lines[5]+'\n')
            f.writelines('工作单位：'+lines[6]+'\n')
            f.writelines('\n')

def writeWord(all_content):
    # 创建文档对象
    document = Document()
    # document.styles['Normal'].font.name = u'宋体'
    # 设置正文中文字体
    microsoft_font = u'微软雅黑'  # u 表示后面的字符串以 Unicode 格式进行编码
    area = qn('w:eastAsia')
    document.styles['Normal'].font.name = microsoft_font
    document.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)

    for lines in all_content:
        if lines[0]:
            # 添加分页符
            if lines[0] != '北京':
                document.add_page_break()
            # 设置文档标题，中文要用unicode字符串
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(u'\n'+lines[0]+'校友会\n')
            run.font.size = Pt(20)
        q = document.add_paragraph()
        q.add_run(u'职务：'+lines[1]+'\n')
        q.add_run(u'姓名：'+lines[2]+'\n')
        q.add_run(u'毕业学院：'+lines[3]+'\n')
        q.add_run(u'毕业时间：'+lines[4]+'\n')
        q.add_run(u'手机：'+lines[5]+'\n')
        q.add_run(u'工作单位：'+lines[6]+'\n')
 
    # 设置文档标题，中文要用unicode字符串
    # document.add_heading(u'我的一个新文档',0)
    
    # 往文档中添加段落
    # p = document.add_paragraph('This is a paragraph having some ')
    # p.add_run('bold ').bold = True
    # p.add_run('and some ')
    # p.add_run('italic.').italic = True
    
    # 添加一级标题
    # document.add_heading(u'一级标题, level = 1',level = 1)
    # document.add_paragraph('Intense quote',style = 'IntenseQuote')
    
    # 添加图片，并指定宽度
    #document.add_picture('e:/docs/pic.png',width = Inches(1.25))
    
    # 保存文档
    document.save('E:\\Yoc\\Course\\2021.docx')

def main():
    # importlib.reload(sys)
    # sys.setdefaultencoding('utf-8')
    all_content = list()
    readExcel(all_content)
    writeTxt(all_content)
    writeWord(all_content)

if __name__ == '__main__':
    main()
